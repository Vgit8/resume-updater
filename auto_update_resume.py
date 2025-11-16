
# auto_update_resume.py
# Script to fetch Naukri profile (or public URL), extract keywords and ATS summary,
# update your DOCX resume automatically, and create backups.

import os
import time
import re
import shutil
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service as ChromeService

NAUKRI_USER = os.environ.get("NAUKRI_USER")
NAUKRI_PASS = os.environ.get("NAUKRI_PASS")
RESUME_PATH = os.environ.get("RESUME_PATH", "Viswanatha_Resume_Updated.docx")
BACKUP_DIR = os.environ.get("BACKUP_DIR", "resume_backups")
NAUKRI_PROFILE_URL = os.environ.get("NAUKRI_PROFILE_URL", "https://www.naukri.com/mnjuser/profile")
CHROMEDRIVER_PATH = os.environ.get("CHROMEDRIVER_PATH", "/usr/bin/chromedriver")
CHROME_BINARY_PATH = os.environ.get("CHROME_BINARY_PATH", "/usr/bin/google-chrome-stable")

def get_driver():
    options = webdriver.ChromeOptions()
    options.add_argument("--headless=new")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    if CHROME_BINARY_PATH and os.path.exists(CHROME_BINARY_PATH):
        options.binary_location = CHROME_BINARY_PATH

    service = ChromeService(executable_path=CHROMEDRIVER_PATH)
    driver = webdriver.Chrome(service=service, options=options)
    driver.set_window_size(1200, 900)
    return driver

def login_and_fetch_html():
    if not NAUKRI_USER or not NAUKRI_PASS:
        import requests
        r = requests.get(NAUKRI_PROFILE_URL, timeout=30)
        r.raise_for_status()
        return r.text

    driver = get_driver()
    try:
        driver.get("https://www.naukri.com/nlogin/login")
        time.sleep(3)

        try:
            email = driver.find_element(By.ID, "username")
            pw = driver.find_element(By.ID, "password")
        except:
            email = driver.find_element(By.NAME, "email")
            pw = driver.find_element(By.NAME, "PASSWORD")

        email.send_keys(NAUKRI_USER)
        pw.send_keys(NAUKRI_PASS)
        pw.send_keys(Keys.RETURN)
        time.sleep(4)

        driver.get(NAUKRI_PROFILE_URL)
        time.sleep(3)

        return driver.page_source
    finally:
        driver.quit()

def extract_summary_and_keywords(html):

    soup = BeautifulSoup(html, "html.parser")
    profile_text = ""

    selectors = [
        ".profileSection", ".profile-desc", ".profileSummary", ".user-details",
        "#summary", ".summary", ".cvHeader", ".professionalSummary"
    ]

    for s in selectors:
        el = soup.select_one(s)
        if el and el.get_text(strip=True):
            profile_text = el.get_text(separator=" ", strip=True)
            break

    if not profile_text:
        paragraphs = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
        profile_text = " ".join(paragraphs)

    txt = profile_text or ""

    keyword_map = {
        r"Medicare": "Medicare",
        r"Medicaid": "Medicaid",
        r"claims?": "Claims Adjudication",
        r"denial": "Denial Management",
        r"high[- ]?dollar": "High-dollar Claims",
        r"medical record": "Medical Record Review",
        r"benefit": "Benefits Verification",
        r"SLA": "SLA Compliance",
        r"Jira": "Jira",
        r"Excel": "MS Excel",
        r"SQL": "SQL"
    }

    found = []
    for pat, canon in keyword_map.items():
        if re.search(pat, txt, re.IGNORECASE):
            found.append(canon)

    if not found:
        found = ["Medicare", "Medicaid", "Claims Adjudication", "Denial Management", "Jira", "Excel", "SQL"]

    keywords = sorted(list(dict.fromkeys(found)))[:12]

    summary = (
        "Healthcare Claims Specialist with 6+ years in Medicare/Medicaid claims adjudication, "
        "denial management, and high-dollar claim review. Skilled in medical record analysis, "
        "benefits verification, and SLA compliance. Proficient in Jira, MS Excel, and SQL."
    )

    return summary, keywords

def backup_and_update_docx(summary, keywords):

    now = datetime.now().strftime("%Y%m%d_%H%M%S")

    if os.path.exists(RESUME_PATH):
        os.makedirs(BACKUP_DIR, exist_ok=True)
        shutil.copy2(RESUME_PATH, os.path.join(BACKUP_DIR, f"backup_{now}.docx"))

    doc = Document(RESUME_PATH) if os.path.exists(RESUME_PATH) else Document()
    paras = doc.paragraphs

    replaced = False
    for i, p in enumerate(paras):
        if p.text.strip().upper().startswith("SUMMARY"):
            paras[i+1].text = summary
            replaced = True
            break

    if not replaced:
        doc.add_heading("SUMMARY", level=1)
        doc.add_paragraph(summary)

    replaced_keys = False
    for i, p in enumerate(paras):
        if p.text.strip().upper().startswith(("SKILLS", "KEYWORDS")):
            paras[i+1].text = ", ".join(keywords)
            replaced_keys = True
            break

    if not replaced_keys:
        doc.add_heading("KEYWORDS", level=2)
        doc.add_paragraph(", ".join(keywords))

    doc.save(RESUME_PATH)

def main():
    html = login_and_fetch_html()
    summary, keywords = extract_summary_and_keywords(html)
    backup_and_update_docx(summary, keywords)

if __name__ == "__main__":
    main()
