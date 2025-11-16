#!/usr/bin/env python3
# auto_update_resume_simple.py
# Simple resume updater (Option A) - no external login.
# - Reads Viswanatha_Resume_Updated.docx in the repo root
# - Extracts keywords from Skills / Experience sections
# - Generates a 3-line ATS-friendly summary and a keywords line
# - Backs up previous resume to resume_backups/ folder

import re, os, shutil
from pathlib import Path
from datetime import datetime
from docx import Document

RESUME = Path("Viswanatha_Resume_Updated.docx")
BACKUP_DIR = Path("resume_backups")
KEYWORDS_LIMIT = 12

def read_text_from_docx(path):
    doc = Document(path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    return texts, doc

def extract_keywords(texts):
    joined = " ".join(texts)
    keywords = set()
    boosters = ["Medicare","Medicaid","Claims","Denial","Adjudication","Jira","Excel","SQL","SLA",
                "Audit","Medical Record","Benefits Verification","Provider","Contract","Denial Management"]
    for b in boosters:
        if re.search(r"\b" + re.escape(b) + r"\b", joined, re.IGNORECASE):
            keywords.add(b)
    # safe regex: place hyphen escaped or at end so it doesn't form ranges
    for w in re.findall(r"\b[A-Z][A-Za-z0-9/\-]{3,}\b", joined):
        if len(keywords) >= KEYWORDS_LIMIT:
            break
        keywords.add(w)
    if not keywords:
        keywords = set(["Medicare","Medicaid","Claims Adjudication","Denial Management","Jira","MS Excel","SQL"])
    return list(keywords)[:KEYWORDS_LIMIT]

def build_summary(keywords):
    summary = ("Healthcare Claims Specialist with 6+ years experience in Medicare/Medicaid claims adjudication, "
               "denial management, and high-dollar claim review. Skilled in medical record analysis, benefits verification, "
               "and SLA compliance. Proficient in {}.").format(", ".join(keywords[:4]))
    return summary

def update_docx(path, summary, keywords):
    doc = Document(path)
    paras = doc.paragraphs
    summary_idx = None
    for i,p in enumerate(paras):
        if p.text.strip().upper().startswith("SUMMARY"):
            summary_idx = i
            break
    if summary_idx is not None and summary_idx+1 < len(paras):
        paras[summary_idx+1].text = summary
    else:
        doc.add_heading("SUMMARY", level=2)
        doc.add_paragraph(summary)
    key_idx = None
    for i,p in enumerate(paras):
        if p.text.strip().upper().startswith("KEY") or p.text.strip().upper().startswith("SKILL"):
            key_idx = i
            break
    kw_line = ", ".join(keywords)
    if key_idx is not None and key_idx+1 < len(paras):
        paras[key_idx+1].text = kw_line
    else:
        doc.add_heading("KEYWORDS", level=2)
        doc.add_paragraph(kw_line)
    doc.save(path)

def backup_and_update(resume_path, summary, keywords):
    if resume_path.exists():
        BACKUP_DIR.mkdir(exist_ok=True)
        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.copy2(resume_path, BACKUP_DIR / f"backup_{now}.docx")
    update_docx(resume_path, summary, keywords)

def main():
    if not RESUME.exists():
        print("ERROR: Resume file not found:", RESUME)
        return 2
    texts, _ = read_text_from_docx(RESUME)
    keywords = extract_keywords(texts)
    summary = build_summary(keywords)
    backup_and_update(RESUME, summary, keywords)
    print("Updated resume:", RESUME)
    print("Keywords:", keywords)
    return 0

if __name__ == '__main__':
    exit(main())
